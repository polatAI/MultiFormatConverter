import PyPDF2
from docx import Document
from docx.shared import Inches
from moviepy.editor import VideoFileClip
from pdf2pptx import Converter
from pdf2pptx.pptx import Presentation
from pptx2pdf import convert as pptx_to_pdf
from weasyprint import HTML
from bs4 import BeautifulSoup
import os

def convert_pdf_to_docx(pdf_file, docx_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    
    doc = Document()
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text = page.extract_text()
        doc.add_paragraph(text)
        
    doc.save(docx_file)
    print(f"{pdf_file} dosyası {docx_file} adında DOCX dosyasına dönüştürüldü.")

def convert_docx_to_pdf(docx_file, pdf_file):
    doc = Document(docx_file)
    pdf = PyPDF2.PdfWriter()

    for para in doc.paragraphs:
        pdf.addPage(PyPDF2.PageObject(para.text))

    with open(pdf_file, "wb") as out_file:
        pdf.write(out_file)
    print(f"{docx_file} dosyası {pdf_file} adında PDF dosyasına dönüştürüldü.")

def convert_mp4_to_mp3(mp4_file, mp3_file):
    video_clip = VideoFileClip(mp4_file)
    audio_clip = video_clip.audio
    audio_clip.write_audiofile(mp3_file)
    audio_clip.close()
    video_clip.close()
    print(f"{mp4_file} dosyası {mp3_file} adında MP3 dosyasına dönüştürüldü.")

def convert_pdf_to_pptx(pdf_file, pptx_file):
    pdf_converter = Converter(pdf_file)
    pdf_converter.convert(pptx_file, start=0, end=None)
    print(f"{pdf_file} dosyası {pptx_file} adında PPTX dosyasına dönüştürüldü.")

def convert_pptx_to_pdf(pptx_file, pdf_file):
    pptx_to_pdf(pptx_file, pdf_file)
    print(f"{pptx_file} dosyası {pdf_file} adında PDF dosyasına dönüştürüldü.")

def convert_html_to_pdf(html_file, pdf_file):
    HTML(string=open(html_file, "r").read()).write_pdf(pdf_file)
    print(f"{html_file} dosyası {pdf_file} adında PDF dosyasına dönüştürüldü.")

def convert_pdf_to_html(pdf_file, html_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()

    soup = BeautifulSoup(f"<pre>{text}</pre>", "html.parser")
    with open(html_file, "w", encoding="utf-8") as file:
        file.write(str(soup))
    print(f"{pdf_file} dosyası {html_file} adında HTML dosyasına dönüştürüldü.")

def main():
    print("1. PDF'den DOCX'e Dönüştürme")
    print("2. DOCX'den PDF'e Dönüştürme")
    print("3. MP4'ten MP3'e Dönüştürme")
    print("4. PDF'den PPTX'e Dönüştürme")
    print("5. PPTX'ten PDF'e Dönüştürme")
    print("6. HTML'den PDF'e Dönüştürme")
    print("7. PDF'den HTML'e Dönüştürme")

    choice = input("Bir seçenek girin (1/2/3/4/5/6/7): ")

    if choice == "1":
        pdf_file = input("PDF dosyasının adını girin: ")
        docx_file = input("DOCX dosyasının adını girin: ")
        convert_pdf_to_docx(pdf_file, docx_file)

    elif choice == "2":
        docx_file = input("DOCX dosyasının adını girin: ")
        pdf_file = input("PDF dosyasının adını girin: ")
        convert_docx_to_pdf(docx_file, pdf_file)

    elif choice == "3":
        mp4_file = input("MP4 dosyasının adını girin: ")
        mp3_file = input("MP3 dosyasının adını girin: ")
        convert_mp4_to_mp3(mp4_file, mp3_file)

    elif choice == "4":
        pdf_file = input("PDF dosyasının adını girin: ")
        pptx_file = input("PPTX dosyasının adını girin: ")
        convert_pdf_to_pptx(pdf_file, pptx_file)
    
    elif choice == "5":  
        pptx_file = input("PPTX dosyasının adını girin: ")
        pdf_file = input("PDF dosyasının adını girin: ")
        convert_pptx_to_pdf(pptx_file, pdf_file)

    elif choice == "6":  
        html_file = input("HTML dosyasının adını girin: ")
        pdf_file = input("PDF dosyasının adını girin: ")
        convert_html_to_pdf(html_file, pdf_file)
        
    elif choice == "7":  
        pdf_file = input("PDF dosyasının adını girin: ")
        html_file = input("HTML dosyasının adını girin: ")
        convert_pdf_to_html(pdf_file, html_file)

    else:
        print("Geçersiz seçenek.")

if __name__ == "__main__":
    main()
